import os
import copy
import random
import tempfile
import zipfile
import re
from io import BytesIO
import openpyxl
from datetime import datetime

from flask import Flask, request, send_file, render_template_string, session, redirect, url_for
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)
app.secret_key = "viphap_exam_mixer_super_secret"

ALLOWED_USERS = {
    "admin": "admin123",
    "viphap": "viphap999"
}

LOGIN_TEMPLATE = """
<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Đăng Nhập - Tool Trộn Đề Thi</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body { font-family: 'Inter', sans-serif; background-color: #0f172a; color: #f8fafc; }
        .glass { background: rgba(30, 41, 59, 0.7); backdrop-filter: blur(10px); border: 1px solid rgba(255,255,255,0.1); }
        .btn-gradient { background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%); transition: transform 0.2s, box-shadow 0.2s; }
        .btn-gradient:hover { transform: translateY(-2px); box-shadow: 0 10px 15px -3px rgba(139, 92, 246, 0.4); }
    </style>
</head>
<body class="min-h-screen flex items-center justify-center p-4">
    <div class="glass w-full max-w-sm rounded-2xl p-8 shadow-2xl">
        <div class="text-center mb-8">
            <h1 class="text-3xl font-bold bg-clip-text text-transparent bg-gradient-to-r from-blue-400 to-purple-500 mb-2">Đăng Nhập</h1>
            <p class="text-slate-400 text-sm">Vui lòng đăng nhập để sử dụng</p>
        </div>
        
        {% if error %}
        <div class="mb-4 p-3 bg-red-500/20 border border-red-500 rounded text-red-400 text-sm text-center">
            {{ error }}
        </div>
        {% endif %}
        
        <form action="/login" method="POST" class="space-y-6">
            <div>
                <label class="block text-sm font-medium text-slate-300 mb-2">Tên đăng nhập:</label>
                <input type="text" name="username" 
                    class="w-full px-4 py-3 bg-slate-800/50 border border-slate-700 rounded-lg focus:outline-none focus:ring-2 focus:ring-blue-500 text-white placeholder-slate-500 transition-colors"
                    required>
            </div>
            
            <div>
                <label class="block text-sm font-medium text-slate-300 mb-2">Mật khẩu:</label>
                <input type="password" name="password" 
                    class="w-full px-4 py-3 bg-slate-800/50 border border-slate-700 rounded-lg focus:outline-none focus:ring-2 focus:ring-blue-500 text-white placeholder-slate-500 transition-colors"
                    required>
            </div>
            
            <button type="submit" class="btn-gradient w-full py-3.5 rounded-lg text-white font-semibold text-lg shadow-lg">
                Đăng Nhập
            </button>
        </form>
    </div>
</body>
</html>
"""

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Tool Trộn Đề Thi Trắc Nghiệm - 2025</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body { font-family: 'Inter', sans-serif; background-color: #0f172a; color: #f8fafc; }
        .glass { background: rgba(30, 41, 59, 0.7); backdrop-filter: blur(10px); border: 1px solid rgba(255,255,255,0.1); }
        .btn-gradient { background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%); transition: transform 0.2s, box-shadow 0.2s; }
        .btn-gradient:hover { transform: translateY(-2px); box-shadow: 0 10px 15px -3px rgba(139, 92, 246, 0.4); }
    </style>
</head>
<body class="min-h-screen flex items-center justify-center p-4">
    <div class="glass w-full max-w-lg rounded-2xl p-8 shadow-2xl">
        <div class="text-center mb-8">
            <h1 class="text-3xl font-bold bg-clip-text text-transparent bg-gradient-to-r from-blue-400 to-purple-500 mb-2">Hệ Thống Trộn Đề Tự Do</h1>
            <p class="text-slate-400 text-sm">Trộn trực tiếp từ văn bản Word - Ép 2 cột</p>
        </div>
        
        <div class="flex justify-between items-center mb-6 px-4 py-3 bg-slate-800/30 rounded-lg border border-slate-700/50">
            <div class="text-sm">Xin chào, <span class="font-bold text-blue-400">{{ username }}</span> 👋</div>
            <a href="/logout" class="text-xs px-3 py-1.5 bg-red-500/20 text-red-400 rounded hover:bg-red-500/40 transition-colors">Đăng xuất</a>
        </div>
        
        <form id="mixForm" action="/mix" method="POST" enctype="multipart/form-data" class="space-y-6">
            <div>
                <label class="block text-sm font-medium text-slate-300 mb-2">Tải lên MÃ ĐỀ GỐC (.docx):</label>
                <input type="file" name="exam_file" accept=".docx" 
                    class="w-full px-4 py-2.5 bg-slate-800/50 border border-slate-700 rounded-lg focus:outline-none focus:ring-2 focus:ring-blue-500 text-slate-300 transition-colors file:mr-4 file:py-2 file:px-4 file:rounded-full file:border-0 file:text-sm file:font-semibold file:bg-blue-500/10 file:text-blue-400 hover:file:bg-blue-500/20"
                    required>
                <p class="text-xs text-slate-400 mt-2">Dùng luôn dạng văn bản. Hãy sửa đề để MỖI PHƯƠNG ÁN NẰM TRÊN 1 DÒNG và bôi vàng đáp án đúng.</p>
            </div>

            <div>
                <label class="block text-sm font-medium text-slate-300 mb-2">Danh sách Mã đề (VD: 301-304 hoặc 101,102):</label>
                <input type="text" name="exam_codes" value="301-304" 
                    class="w-full px-4 py-3 bg-slate-800/50 border border-slate-700 rounded-lg focus:outline-none focus:ring-2 focus:ring-blue-500 text-white placeholder-slate-500 transition-colors"
                    required>
            </div>
            
            <div>
                <label class="block text-sm font-medium text-slate-300 mb-2">Cỡ chữ (Font Size):</label>
                <select name="font_size" class="w-full px-4 py-3 bg-slate-800/50 border border-slate-700 rounded-lg focus:outline-none focus:ring-2 focus:ring-blue-500 text-white transition-colors">
                    <option value="12">Size 12 (Nhỏ gọn)</option>
                    <option value="13" selected>Size 13 (Tiêu chuẩn)</option>
                    <option value="14">Size 14 (To rõ ràng)</option>
                </select>
            </div>
            
            <button type="submit" id="submitBtn" class="btn-gradient w-full py-3.5 rounded-lg text-white font-semibold text-lg shadow-lg flex justify-center items-center">
                <span>Bắt Đầu Trộn Ngay</span>
                <svg class="w-5 h-5 ml-2" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M13 10V3L4 14h7v7l9-11h-7z"></path></svg>
            </button>
        </form>
        
        <div id="loading" class="hidden mt-6 text-center">
            <div class="inline-block animate-spin w-8 h-8 rounded-full border-4 border-blue-500 border-t-transparent mb-2"></div>
            <p class="text-sm text-slate-400 animate-pulse">Hệ thống đang xử lý và căn lều tiết kiệm giấy...</p>
        </div>
    </div>

    <script>
        document.getElementById('mixForm').addEventListener('submit', function() {
            document.getElementById('submitBtn').classList.add('hidden');
            document.getElementById('loading').classList.remove('hidden');
            setTimeout(() => {
                document.getElementById('submitBtn').classList.remove('hidden');
                document.getElementById('loading').classList.add('hidden');
            }, 3000);
        });
    </script>
</body>
</html>
"""

def has_answer_mark(element):
    from docx.oxml.ns import qn
    for highlight in element.xpath('.//w:highlight'):
        return True
    for shd in element.xpath('.//w:shd'):
        fill = shd.get(qn('w:fill'))
        if fill and fill != 'auto' and fill.lower() not in ['000000', 'ffffff', 'none']:
            return True
    for color in element.xpath('.//w:color'):
        val = color.get(qn('w:val'))
        if val and val != 'auto' and val.lower() not in ['000000']:
            return True
    for u in element.xpath('.//w:u'):
        val = u.get(qn('w:val'))
        if val and val != 'none':
            return True
    return False

def remove_highlights(element):
    for highlight in element.xpath('.//w:highlight'):
        highlight.getparent().remove(highlight)
    for shd in element.xpath('.//w:shd'):
        shd.getparent().remove(shd)
    for u in element.xpath('.//w:u'):
        u.getparent().remove(u)
    for color in element.xpath('.//w:color'):
        color.getparent().remove(color)

def parse_freeform_docx(filename):
    doc = Document(filename)
    questions = []
    current_q = None
    current_part = 1
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text: continue
            
        if "HẾT" in text.upper():
            clean = text.upper().replace("-", "").replace("_", "").replace("*", "").replace("—", "").replace("–", "").replace(".", "").strip()
            if clean == "HẾT" or clean.startswith("HẾT"):
                break
            
        part_match = re.search(r'PHẦN ([I]+)', text)
        if part_match:
            r = part_match.group(1)
            if r == 'I': current_part = 1
            elif r == 'II': current_part = 2
            elif r == 'III': current_part = 3
            continue
            
        câu_match = re.match(r'^Câu\s+\d+[:.]', text, re.IGNORECASE)
        if câu_match:
            if current_q and (len(current_q['opts']) > 0 or current_q['part'] == 3):
                questions.append(current_q)
            current_q = {'part': current_part, 'head': [p], 'opts': [], 'ans': None}
            continue
            
        if current_q:
            if current_q['part'] == 3:
                ans_match = re.match(r'^(ĐÁP ÁN|Đáp án)[\s:]*(.*)', text, re.IGNORECASE)
                if ans_match:
                    val = ans_match.group(2).strip()
                    if val:
                        current_q['ans'] = val
                    continue
                elif has_answer_mark(p._p):
                    ans_text = text.strip()
                    # Clean trailing dot if it's like a standalone number list item e.g. "2."
                    if ans_text and ans_text[-1] == '.' and ans_text[:-1].isdigit():
                        ans_text = ans_text[:-1]
                    current_q['ans'] = ans_text
                    continue
                else:
                    current_q['head'].append(p)
                continue

            opt_match = re.match(r'^([A-D]|[a-d])[.)]\s+(.*)', text)
            if opt_match:
                opt_char = opt_match.group(1).upper()
                is_correct = has_answer_mark(p._p)
                current_q['opts'].append(p)
                if is_correct:
                    if current_q['part'] == 1:
                        current_q['ans'] = opt_char
            else:
                if len(current_q['opts']) == 0:
                    current_q['head'].append(p)

    if current_q and (len(current_q['opts']) > 0 or current_q['part'] == 3):
        questions.append(current_q)
        
    for q in questions:
        if q['part'] == 2 and not q['ans']:
            ans_str = []
            for opt_p in q['opts']:
                if has_answer_mark(opt_p._p):
                    ans_str.append('Đ')
                else:
                    ans_str.append('S')
            q['ans'] = "".join(ans_str)
            if len(q['ans']) < 4:
                q['ans'] = q['ans'].ljust(4, 'S')
                
        if not q['ans']: q['ans'] = "?" # fallback
            
    return questions

def get_or_create_r(text, is_bold=True):
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    if is_bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    r.append(t)
    return r

def strip_option_prefix(new_p):
    t_nodes = new_p.xpath('.//w:t')
    full_text = "".join(t.text for t in t_nodes if t.text)
    match = re.match(r'^\s*([A-D]|[a-d])[.)]\s*', full_text)
    if match:
        to_remove = len(match.group(0))
        for t in t_nodes:
            if not t.text: continue
            if to_remove >= len(t.text):
                to_remove -= len(t.text)
                t.text = ""
            elif to_remove > 0:
                t.text = t.text[to_remove:]
                to_remove = 0
            else:
                break
                
def strip_question_prefix(new_p):
    t_nodes = new_p.xpath('.//w:t')
    full_text = "".join(t.text for t in t_nodes if t.text)
    match = re.match(r'^\s*Câu\s+\d+[:.]\s*', full_text, flags=re.IGNORECASE)
    if match:
        to_remove = len(match.group(0))
        for t in t_nodes:
            if not t.text: continue
            if to_remove >= len(t.text):
                to_remove -= len(t.text)
                t.text = ""
            elif to_remove > 0:
                t.text = t.text[to_remove:]
                to_remove = 0
            else:
                break

def insert_p(doc, p):
    sectPrs = doc._body._body.xpath('./w:sectPr')
    if sectPrs:
        sectPrs[0].addprevious(p)
    else:
        doc._body._body.append(p)

def normalize_p(p_node, font_name='Times New Roman', font_size_val='26'):
    p_style = p_node.xpath('./w:pPr/w:pStyle')
    if p_style:
        p_style[0].set(qn('w:val'), 'Normal')
        
    for r in p_node.xpath('.//w:r'):
        rPrs = r.xpath('./w:rPr')
        if not rPrs:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
            rPrs = [rPr]
        rPr = rPrs[0]
        
        rFonts = rPr.xpath('./w:rFonts')
        if not rFonts:
            rFont = OxmlElement('w:rFonts')
            rPr.append(rFont)
            rFonts = [rFont]
        for rf in rFonts:
            rf.set(qn('w:ascii'), font_name)
            rf.set(qn('w:hAnsi'), font_name)
            rf.set(qn('w:cs'), font_name)
            
        sz_nodes = rPr.xpath('./w:sz')
        if not sz_nodes:
            sz_node = OxmlElement('w:sz')
            rPr.append(sz_node)
            sz_nodes = [sz_node]
        for s in sz_nodes:
            s.set(qn('w:val'), font_size_val)
            
        szCs_nodes = rPr.xpath('./w:szCs')
        if not szCs_nodes:
            szCs_node = OxmlElement('w:szCs')
            rPr.append(szCs_node)
            szCs_nodes = [szCs_node]
        for s in szCs_nodes:
            s.set(qn('w:val'), font_size_val)

def replace_exam_code(doc, new_code):
    all_paras = list(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_paras.extend(c.paragraphs)
                
    for section in doc.sections:
        all_paras.extend(section.header.paragraphs)
        for t in section.header.tables:
            for row in t.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)
        all_paras.extend(section.footer.paragraphs)
        for t in section.footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)
                
    for p in all_paras:
        match = re.search(r'(Mã đề(?: thi)?[\s:]*)(\d+)', p.text, re.IGNORECASE)
        if match:
            start_idx = match.start(2)
            end_idx = match.end(2)
            
            current_idx = 0
            new_code_inserted = False
            
            for r in p.runs:
                if not r.text: continue
                r_start = current_idx
                r_end = current_idx + len(r.text)
                
                if r_end > start_idx and r_start < end_idx:
                    keep_before = ""
                    if r_start < start_idx:
                        keep_before = r.text[:start_idx - r_start]
                        
                    keep_after = ""
                    if r_end > end_idx:
                        keep_after = r.text[end_idx - r_start:]
                        
                    if not new_code_inserted:
                        r.text = keep_before + new_code + keep_after
                        new_code_inserted = True
                    else:
                        r.text = keep_before + keep_after
                
                current_idx = r_end

def generate_exam_linear(questions, exam_code, original_path, font_size_val='26'):
    doc = Document(original_path)
    replace_exam_code(doc, str(exam_code))
    # We will use the original doc as template to keep headers right.
    # We clear all paragraphs that look like questions or parts
    # Actually, simpler: find the first "PHẦN I" and delete everything from there to the end!
    found_start = False
    for p in doc.paragraphs:
        if not found_start:
            if "PHẦN I" in p.text.upper():
                found_start = True
                p._element.getparent().remove(p._element)
        else:
            p._element.getparent().remove(p._element)
            
    # Also remove tables if any (except header tables which are usually table 0 if they exist BEFORE PHẦN I)
    # Be safe: don't touch tables unless they are obviously inside questions. Usually no tables.
    
    part_1 = [copy.copy(q) for q in questions if q['part'] == 1]
    part_2 = [copy.copy(q) for q in questions if q['part'] == 2]
    part_3 = [copy.copy(q) for q in questions if q['part'] == 3]
    
    random.shuffle(part_1)
    random.shuffle(part_2)
    random.shuffle(part_3)
    
    answers = []
    
    # Process part 1
    if part_1:
        p_hdr = doc.add_paragraph("PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn (3,0 điểm). Thí sinh trả lời từ câu 1 đến câu 12. Mỗi câu hỏi thí sinh chỉ chọn một phương án.")
        normalize_p(p_hdr._p, font_size_val=font_size_val)
        for i, q in enumerate(part_1):
            q_num = i + 1
            # Append headers
            for j, p in enumerate(q['head']):
                new_p = copy.deepcopy(p._p)
                remove_highlights(new_p)
                if j == 0:
                    strip_question_prefix(new_p)
                    new_p.insert(0, get_or_create_r(f"Câu {q_num}: ", True))
                normalize_p(new_p, font_size_val=font_size_val)
                insert_p(doc, new_p)
            
            opts = copy.copy(q['opts'])
            ans_char = q['ans']
            ans_idx = ord(ans_char) - ord('A') if len(ans_char)==1 and ans_char in "ABCD" else 0
            
            opt_pairs = list(enumerate(opts))
            random.shuffle(opt_pairs)
            
            # Create a 2x2 table for options
            table = doc.add_table(rows=2, cols=2)
            
            new_ans_idx = -1
            opt_labels = ["A.", "B.", "C.", "D."]
            for new_idx, (old_idx, opt_p) in enumerate(opt_pairs):
                if old_idx == ans_idx: new_ans_idx = new_idx
                
                r_idx = new_idx // 2
                c_idx = new_idx % 2
                cell = table.cell(r_idx, c_idx)
                for p in cell.paragraphs: p._element.getparent().remove(p._element)
                
                new_p = copy.deepcopy(opt_p._p)
                remove_highlights(new_p)
                strip_option_prefix(new_p)
                new_p.insert(0, get_or_create_r(opt_labels[new_idx] + " ", False))
                normalize_p(new_p, font_size_val=font_size_val)
                cell._tc.append(new_p)
                
            final_answer = chr(ord('A') + new_ans_idx)
            answers.append({'part': 1, 'q_num': q_num, 'ans': final_answer})
            
    # Process part 2
    if part_2:
        p_hdr = doc.add_paragraph("PHẦN II. Câu trắc nghiệm đúng sai (2,0 điểm). Thí sinh trả lời từ câu 1 đến câu 2. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai.")
        normalize_p(p_hdr._p, font_size_val=font_size_val)
        for i, q in enumerate(part_2):
            q_num = i + 1
            for j, p in enumerate(q['head']):
                new_p = copy.deepcopy(p._p)
                remove_highlights(new_p)
                if j == 0:
                    strip_question_prefix(new_p)
                    new_p.insert(0, get_or_create_r(f"Câu {q_num}: ", True))
                normalize_p(new_p, font_size_val=font_size_val)
                insert_p(doc, new_p)
            
            opts = copy.copy(q['opts'])
            ans_str = q['ans']
            
            opt_pairs = list(enumerate(opts))
            random.shuffle(opt_pairs)
            
            new_ans_chars = [''] * 4
            opt_labels = ["a)", "b)", "c)", "d)"]
            
            # Use 1x4 lines because T/F options are usually long sentences
            for new_idx, (old_idx, opt_p) in enumerate(opt_pairs):
                new_ans_chars[new_idx] = ans_str[old_idx] if old_idx < len(ans_str) else 'S'
                new_p = copy.deepcopy(opt_p._p)
                remove_highlights(new_p)
                strip_option_prefix(new_p)
                new_p.insert(0, get_or_create_r(opt_labels[new_idx] + " ", False))
                normalize_p(new_p, font_size_val=font_size_val)
                insert_p(doc, new_p)
                
            final_answer = "".join(new_ans_chars)
            answers.append({'part': 2, 'q_num': q_num, 'ans': final_answer})
            
    # Process part 3
    if part_3:
        p_hdr = doc.add_paragraph("PHẦN III. Câu trắc nghiệm trả lời ngắn (2,0 điểm). Thí sinh trả lời từ câu 1 đến câu 8.")
        normalize_p(p_hdr._p, font_size_val=font_size_val)
        for i, q in enumerate(part_3):
            q_num = i + 1
            for j, p in enumerate(q['head']):
                new_p = copy.deepcopy(p._p)
                remove_highlights(new_p)
                if j == 0:
                    strip_question_prefix(new_p)
                    new_p.insert(0, get_or_create_r(f"Câu {q_num}: ", True))
                normalize_p(new_p, font_size_val=font_size_val)
                insert_p(doc, new_p)
            answers.append({'part': 3, 'q_num': q_num, 'ans': q['ans']})
    
    # HET
    p_end = doc.add_paragraph("-------------- HẾT ---------------")
    p_end.alignment = 1 # Center
    
    # Save to memory
    io_stream = BytesIO()
    doc.save(io_stream)
    io_stream.seek(0)
    return io_stream.getvalue(), answers


@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        
        if username in ALLOWED_USERS and ALLOWED_USERS[username] == password:
            session['user'] = username
            return redirect(url_for('index'))
        else:
            error = "Tài khoản hoặc mật khẩu không chính xác!"
            
    return render_template_string(LOGIN_TEMPLATE, error=error)

@app.route("/logout")
def logout():
    session.pop('user', None)
    return redirect(url_for('login'))

@app.route("/")
def index():
    if 'user' not in session:
        return redirect(url_for('login'))
    return render_template_string(HTML_TEMPLATE, username=session['user'])

@app.route("/mix", methods=["POST"])
def mix_exams():
    if 'user' not in session:
        return redirect(url_for('login'))
        
    temp_path = None
    try:
        exam_codes_str = request.form.get("exam_codes", "301-304")
        font_size_pt = request.form.get("font_size", "13")
        font_size_val = str(int(font_size_pt) * 2)
        
        # Parse exam codes
        exam_codes = []
        for part in exam_codes_str.split(','):
            part = part.strip()
            if not part: continue
            if '-' in part:
                try:
                    start_str, end_str = part.split('-')
                    start_num = int(start_str.strip())
                    end_num = int(end_str.strip())
                    padding = len(start_str.strip())
                    for n in range(start_num, end_num + 1):
                        exam_codes.append(str(n).zfill(padding))
                except:
                    exam_codes.append(part)
            else:
                exam_codes.append(part)
        
        if not exam_codes:
            exam_codes = ["101"]
        
        file = request.files.get("exam_file")
        if not file or file.filename == "":
            return "Vui lòng chọn file MÃ ĐỀ GỐC (.docx).", 400
            
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"exam_{random.randint(1000, 9999)}.docx")
        
        file.save(temp_path)
            
        questions = parse_freeform_docx(temp_path)
        if not questions:
            return "Không đọc được câu hỏi nào từ file gốc. Mời kiểm tra lại định dạng.", 400
            
        zip_buffer = BytesIO()
        all_exam_answers = {}
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for code in exam_codes:
                exam_bytes, exam_ans = generate_exam_linear(questions, code, temp_path, font_size_val)
                zip_file.writestr(f"De_{code}.docx", exam_bytes)
                all_exam_answers[code] = exam_ans
                
            # Create Answer Key Excel
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Dap_An"
            
            ws.cell(row=1, column=1, value="Mã Đề")
            first_code = exam_codes[0]
            example_ans = all_exam_answers[first_code]
            for col_idx, ans_info in enumerate(example_ans):
                label = f"Câu {ans_info['part']}.{ans_info['q_num']}"
                ws.cell(row=1, column=col_idx + 2, value=label)
                
            for row_idx, (code, code_answers) in enumerate(all_exam_answers.items()):
                ws.cell(row=row_idx + 2, column=1, value=code)
                for col_idx, ans_info in enumerate(code_answers):
                    ws.cell(row=row_idx + 2, column=col_idx + 2, value=ans_info['ans'])
                    
            excel_buffer = BytesIO()
            wb.save(excel_buffer)
            excel_buffer.seek(0)
            zip_file.writestr("Bang_Dap_An.xlsx", excel_buffer.getvalue())

        zip_buffer.seek(0)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return send_file(
            zip_buffer,
            mimetype="application/zip",
            as_attachment=True,
            download_name=f"TronDe_{timestamp}.zip"
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"Lỗi xảy ra trong quá trình trộn: {str(e)}", 500
    finally:
        if temp_path and os.path.exists(temp_path):
            try: os.remove(temp_path)
            except: pass

if __name__ == "__main__":
    app.run(debug=True, port=5000)
