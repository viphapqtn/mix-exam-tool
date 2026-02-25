from docx import Document
doc = Document("De goc.docx")

shapes = 0
maths = 0

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if '<w:drawing' in p._p.xml or '<v:shape' in p._p.xml:
                    shapes += 1
                if '<m:oMath' in p._p.xml:
                    maths += 1

print("Shapes:", shapes)
print("Math equations:", maths)
