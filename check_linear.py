from docx import Document
doc = Document("test_out_linear.docx")
lines = [p.text for p in doc.paragraphs if p.text.strip()]
with open("check_linear_full.txt", "w", encoding="utf-8") as f:
    f.write("\n=== PARAS ===\n")
    for i, l in enumerate(lines): f.write(f"{i}: {l}\n")
    f.write("\n=== TABLES ===\n")
    if len(doc.tables) > 0:
        for i, row in enumerate(doc.tables[0].rows):
            cells = [c.text.strip() for c in row.cells]
            f.write(f"Table 0 Row {i}: {cells}\n")
