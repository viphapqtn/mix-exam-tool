import json
from docx import Document
doc = Document("De goc.docx")

data = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "cells": []}

for i, table in enumerate(doc.tables):
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                data["cells"].append({"table": i, "row": r, "col": c, "text": text[:100]})

with open("out.json", "w", encoding="utf-8") as f:
    json.dump(data, f, ensure_ascii=False, indent=2)
