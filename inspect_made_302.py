from docx import Document
import json

try:
    doc = Document("Made 302.docx")
    out = {
        "paragraphs": [p.text for p in doc.paragraphs[:10]],
        "tables": len(doc.tables),
        "table_0_rows": []
    }
    if len(doc.tables) > 0:
        table = doc.tables[0]
        for i, row in enumerate(table.rows[:20]):
            cells = [c.text.strip()[:20] for c in row.cells]
            out["table_0_rows"].append(cells)
    
    with open("made_302_out.json", "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print("Done")
except Exception as e:
    print("Error:", e)
