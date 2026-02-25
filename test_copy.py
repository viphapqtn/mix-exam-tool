import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def get_or_create_r(text):
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    # make it bold
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    r.append(rPr)
    r.append(t)
    return r

doc_in = Document("De goc.docx")
table = doc_in.tables[0]

doc_out = Document("Made 302.docx")
for p in doc_out.paragraphs:
    p._element.getparent().remove(p._element)

# Question
q_cell = table.rows[0].cells[1]
for i, p in enumerate(q_cell.paragraphs):
    new_p = copy.deepcopy(p._p)
    if i == 0:
        r = get_or_create_r("Câu X: ")
        # insert at beginning
        new_p.insert(0, r)
    doc_out._body._body.append(new_p)

# Options
opts = ["A. ", "B. ", "C. ", "D. "]
for i in range(4):
    opt_cell = table.rows[i+1].cells[1]
    for j, p in enumerate(opt_cell.paragraphs):
        new_p = copy.deepcopy(p._p)
        if j == 0:
            r = get_or_create_r(opts[i])
            new_p.insert(0, r)
        doc_out._body._body.append(new_p)

doc_out.save("test_linear.docx")
print("Done")
