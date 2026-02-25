from docx import Document

doc = Document("Made 302.docx")
paras = list(doc.paragraphs)
found_start = False
for p in paras:
    if not found_start:
        if "PHẦN I" in p.text.upper():
            found_start = True
            p._element.getparent().remove(p._element)
    else:
        p._element.getparent().remove(p._element)

doc.save("test_del_2.docx")

doc2 = Document("test_del_2.docx")
print("Remaining paragraphs:")
for i, p in enumerate(doc2.paragraphs):
    print(f"Remain {i}: {p.text}")
