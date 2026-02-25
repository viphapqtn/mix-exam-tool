import sys
import codecs
sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())

import re
from docx import Document

doc = Document()
p = doc.add_paragraph()
p.add_run("Mã ")
p.add_run("đề ")
p.add_run("thi ")
p.add_run("30")
p.add_run("2")
p.add_run(" - ")
p.add_run("Trang 1/3")

print("Original runs:", [r.text for r in p.runs])

match = re.search(r'(Mã đề(?: thi)?[\s:]*)(\d+)', p.text, re.IGNORECASE)
if match:
    prefix = match.group(1)
    original_code = match.group(2)
    start_idx = match.start(2)
    end_idx = match.end(2)
    new_code = "101"
    
    current_idx = 0
    new_code_inserted = False
    
    for r in p.runs:
        if not r.text: continue
        r_start = current_idx
        r_end = current_idx + len(r.text)
        
        if r_end > start_idx and r_start < end_idx:
            keep_before = ""
            if r_start < start_idx:
                keep_before = r.text[:start_idx - r_start]
                
            keep_after = ""
            if r_end > end_idx:
                keep_after = r.text[end_idx - r_start:]
                
            if not new_code_inserted:
                r.text = keep_before + new_code + keep_after
                new_code_inserted = True
            else:
                r.text = keep_before + keep_after
        
        current_idx = r_end

print("Replaced runs:", [r.text for r in p.runs])
print("Final text:", p.text)
