from docx import Document

doc = Document("test_replace_exam_code.docx")

with open("test_format_out.txt", "w", encoding="utf-8") as f:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "Mã đề" in p.text or "Mã Đề" in p.text:
                        f.write(f"Found paragraph: {p.text}\n")
                        for r in p.runs:
                            f.write(f"Run text: '{r.text}' - Bold: {r.bold}\n")
