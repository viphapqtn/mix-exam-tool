from docx import Document
import re

doc = Document("Made 302.docx")
new_code = "999"

def process_paras(paras):
    for p in paras:
        if re.search(r'(Mã đề(?: thi)?[\s:]*)\d+', p.text, re.IGNORECASE):
            new_text = re.sub(r'(Mã đề(?: thi)?[\s:]*)\d+', r'\g<1>' + new_code, p.text, flags=re.IGNORECASE)
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)

all_paras = []
for section in doc.sections:
    # Header
    all_paras.extend(section.header.paragraphs)
    for t in section.header.tables:
        for row in t.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    
    # Footer
    all_paras.extend(section.footer.paragraphs)
    for t in section.footer.tables:
        for row in t.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

process_paras(all_paras)

doc.save("test_hf_replace.docx")
print("Done")
