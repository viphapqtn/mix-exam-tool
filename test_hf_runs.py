from docx import Document
import re

doc = Document("Made 302.docx")
new_code = "101"

all_paras = list(doc.paragraphs)
for section in doc.sections:
    all_paras.extend(section.header.paragraphs)
    all_paras.extend(section.footer.paragraphs)
    for t in section.header.tables:
        for row in t.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    for t in section.footer.tables:
        for row in t.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

for p in all_paras:
    match = re.search(r'Mã đề(?: thi)?[\s:]*(\d+)', p.text, re.IGNORECASE)
    if match:
        original_code = match.group(1)
        # Now replace this original_code in the runs of this paragraph
        for r in p.runs:
            if original_code in r.text:
                r.text = r.text.replace(original_code, new_code)

doc.save("test_hf_runs.docx")
print("Done")
