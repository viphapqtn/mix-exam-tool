import copy
from docx import Document

doc = Document("De goc.docx")
table = doc.tables[0]
opts = [table.rows[1], table.rows[2], table.rows[3], table.rows[4]]

new_doc = Document()
new_table = new_doc.add_table(rows=0, cols=2)
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

for opt in opts:
    new_opt_tr = copy.deepcopy(opt._tr)
    for highlight in new_opt_tr.xpath('.//w:highlight', namespaces=nsmap):
        highlight.getparent().remove(highlight)
    new_table._tbl.append(new_opt_tr)

new_doc.save("test_clear_highlight.docx")
print("Saved to test_clear_highlight.docx")
