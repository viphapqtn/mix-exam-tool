import re
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def has_answer_mark(paragraph):
    for run in paragraph.runs:
        if getattr(run.font, 'highlight_color', None):
            return True
        if run.underline:
            return True
    return False

def remove_highlights(element):
    for highlight in element.xpath('.//w:highlight'):
        highlight.getparent().remove(highlight)

def parse_freeform_docx(filename):
    doc = Document(filename)
    questions = []
    current_q = None
    current_part = 1
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        part_match = re.search(r'PHẦN ([I]+)', text)
        if part_match:
            r = part_match.group(1)
            if r == 'I': current_part = 1
            elif r == 'II': current_part = 2
            elif r == 'III': current_part = 3
            continue
            
        câu_match = re.match(r'^Câu\s+\d+[:.]', text, re.IGNORECASE)
        if câu_match:
            if current_q and (len(current_q['opts']) > 0 or current_q['part'] == 3):
                questions.append(current_q)
            current_q = {
                'part': current_part,
                'head': [p],
                'opts': [],
                'ans': None
            }
            continue
            
        if current_q:
            if current_q['part'] == 3:
                ans_match = re.match(r'^(ĐÁP ÁN|Đáp án)[\s:]*(.*)', text, re.IGNORECASE)
                if ans_match:
                    current_q['ans'] = ans_match.group(2).strip()
                else:
                    current_q['head'].append(p)
                continue

            opt_match = re.match(r'^([A-D]|[a-d])[.)]\s+(.*)', text)
            if opt_match:
                opt_char = opt_match.group(1).upper()
                is_correct = has_answer_mark(p)
                current_q['opts'].append({
                    'char': opt_char,
                    'paragraph': p,
                    'is_correct': is_correct
                })
            else:
                if len(current_q['opts']) == 0:
                    current_q['head'].append(p)
                else:
                    current_q['opts'][-1]['paragraph']._p.append(copy.deepcopy(p._p)) # append to previous option instead if lines are broken

    if current_q and (len(current_q['opts']) > 0 or current_q['part'] == 3):
        questions.append(current_q)
        
    return questions

def get_or_create_r(text, is_bold=True):
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    if is_bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    r.append(t)
    return r

def append_paragraph_modified(doc_out, p_in, prefix, is_bold_prefix=True, target_cell=None):
    new_p = copy.deepcopy(p_in._p)
    remove_highlights(new_p)
    # Remove the original A., B., C., D. prefix from the text run
    # This is tricky because the prefix A. could be in the first run.
    # We can try to replace the text in the first run if it matches.
    for run in new_p.xpath('.//w:r'):
        t_nodes = run.xpath('.//w:t')
        if t_nodes:
            t = t_nodes[0]
            val = t.text or ""
            val = val.lstrip()
            rem = re.sub(r'^([A-D]|[a-d])[.)]\s*', '', val)
            if val != rem:
                t.text = rem
                break # only replace first occurrence

    # add new prefix
    r = get_or_create_r(prefix + " ", is_bold_prefix)
    new_p.insert(0, r)
    
    if target_cell is not None:
        target_cell._tc.append(new_p)
    else:
        doc_out._body._body.append(new_p)

# test generate 2x2
doc = Document()
doc.add_paragraph("Test")
# mock q
class MockP:
    pass
# Actually we will just test the modify loop in app.py directly
