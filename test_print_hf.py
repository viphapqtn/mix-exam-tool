from docx import Document

doc = Document("test_hf_runs.docx")
for section in doc.sections:
    for p in section.footer.paragraphs:
        print("Footer para:", p.text.encode('utf-8'))
    for t in section.footer.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    print("Footer table cell para:", p.text.encode('utf-8'))
