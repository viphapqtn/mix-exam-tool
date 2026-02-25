from docx import Document
import re

doc = Document("Made 302.docx")
exam_code = "999"

with open("test_replace_code_out.txt", "w", encoding="utf-8") as f:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "Mã đề" in p.text or "Mã Đề" in p.text:
                        f.write(f"Found paragraph: {p.text}\n")
                        for r in p.runs:
                            f.write(f"Run text: '{r.text}'\n")
                            if re.search(r'Mã đề thi\s+\d+', r.text, re.IGNORECASE):
                                r.text = re.sub(r'Mã đề thi\s+\d+', f'Mã đề thi {exam_code}', r.text, flags=re.IGNORECASE)
                            elif "302" in r.text: # Hardcoded check for 302 if it's separated into its own run
                                r.text = r.text.replace("302", exam_code)
                            elif re.search(r'Mã đề\s+\d+', r.text, re.IGNORECASE):
                                r.text = re.sub(r'Mã đề\s+\d+', f'Mã đề {exam_code}', r.text, flags=re.IGNORECASE)

doc.save("test_replace_code.docx")
print("Done")
