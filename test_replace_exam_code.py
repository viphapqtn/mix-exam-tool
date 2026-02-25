import re
from docx import Document

def replace_exam_code(doc, new_code):
    all_paras = list(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_paras.extend(c.paragraphs)
                
    for p in all_paras:
        if re.search(r'(Mã đề(?: thi)?[\s:]*)\d+', p.text, re.IGNORECASE):
            # preserve original text, just swap out the numbers
            new_text = re.sub(r'(Mã đề(?: thi)?[\s:]*)\d+', r'\g<1>' + new_code, p.text, flags=re.IGNORECASE)
            # Try to preserve styling of the first run
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)

doc = Document("Made 302.docx")
replace_exam_code(doc, "987")
doc.save("test_replace_exam_code.docx")
print("Done")
