from docx import Document
import copy
from docx.oxml import OxmlElement

def insert_p(doc, p):
    sectPrs = doc._body._body.xpath('./w:sectPr')
    if sectPrs:
        sectPrs[0].addprevious(p)
    else:
        doc._body._body.append(p)

doc = Document()
doc.add_paragraph("First")
p = OxmlElement('w:p')
r = OxmlElement('w:r')
t = OxmlElement('w:t')
t.text = "Appended"
r.append(t)
p.append(r)
insert_p(doc, p)
doc.add_paragraph("Second")
doc.save("test_sect_pr.docx")

doc2 = Document("test_sect_pr.docx")
print([p.text for p in doc2.paragraphs])
