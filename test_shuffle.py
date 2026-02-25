import copy
from docx import Document

doc = Document("De goc.docx")
table = doc.tables[0]
rows = table.rows

total_rows = len(rows)
print("Total rows:", total_rows)

# Assuming 6 rows per question
qs = []
for i in range(0, total_rows, 6):
    if i + 5 < total_rows:
        q_rows = rows[i:i+6]
        # Get correct answer
        ans_text = q_rows[5].cells[1].text.strip()
        qs.append({
            'head': q_rows[0],
            'opts': [q_rows[1], q_rows[2], q_rows[3], q_rows[4]],
            'ans_row': q_rows[5],
            'ans': ans_text
        })

print("First question ans:", qs[0]['ans'])

# Build new document by manipulating XML
new_doc = Document()
new_table = new_doc.add_table(rows=0, cols=2)
new_table.style = table.style

for q in qs[:2]:
    new_table._tbl.append(copy.deepcopy(q['head']._tr))
    new_table._tbl.append(copy.deepcopy(q['opts'][0]._tr))
    new_table._tbl.append(copy.deepcopy(q['opts'][1]._tr))
    new_table._tbl.append(copy.deepcopy(q['opts'][2]._tr))
    new_table._tbl.append(copy.deepcopy(q['opts'][3]._tr))
    new_table._tbl.append(copy.deepcopy(q['ans_row']._tr))

new_doc.save("test_shuffle.docx")
print("Saved to test_shuffle.docx")
