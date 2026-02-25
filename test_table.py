import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def get_or_create_r(text, is_bold=True):
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    if is_bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    r.append(t)
    return r

def remove_highlights(element):
    for highlight in element.xpath('.//w:highlight'):
        highlight.getparent().remove(highlight)

doc = Document()
doc.add_paragraph("Câu 1: What is 1+1?")

table = doc.add_table(rows=2, cols=2)
# Optionally remove borders (in python-docx this is done via styles or xml, but default table might have borders - we can use 'Normal Table' which usually has no borders, or we can just hope default has no borders. Actually 'Table Grid' has borders, default might not.)
# Let's see what styles are available
# We can just construct a table.

cell = table.cell(0, 0)
p = cell.paragraphs[0]
r = get_or_create_r("A. 1", False)
p._p.append(r)

cell = table.cell(0, 1)
p = cell.paragraphs[0]
r = get_or_create_r("B. 2", False)
p._p.append(r)

cell = table.cell(1, 0)
p = cell.paragraphs[0]
r = get_or_create_r("C. 3", False)
p._p.append(r)

cell = table.cell(1, 1)
p = cell.paragraphs[0]
r = get_or_create_r("D. 4", False)
p._p.append(r)

doc.save("test_table_options.docx")
print("Done")
